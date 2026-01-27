import sys
from pathlib import Path

# Add project root to path
sys.path.append(str(Path(__file__).parent.parent))

from docx import Document
from backend.core.latex_converter import convert_latex_in_document


def test_latex_conversion():
    # Create a test document
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("Here is an inline formula: $E=mc^2$ inside text.")
    doc.add_paragraph("And here is a display formula:")
    doc.add_paragraph("$$\\frac{a}{b} = c$$")

    # Run conversion
    print("Running conversion...")
    fixes = convert_latex_in_document(doc)

    print(f"Applied {len(fixes)} fixes:")
    for fix in fixes:
        print(f" - {fix['description']}")

    # Check if OMML was inserted
    # We check if the paragraph xml contains 'm:oMath'

    has_omml = False
    for p in doc.paragraphs:
        if "m:oMath" in p._p.xml:
            has_omml = True
            print(f"Found OMML in paragraph: {p.text[:20]}...")

    if has_omml:
        print("SUCCESS: OMML elements found in document.")
        # Save for manual inspection if needed
        output_path = Path("tests/output_latex_test.docx")
        doc.save(output_path)
        print(f"Saved test output to {output_path}")
    else:
        print("FAILURE: No OMML elements found.")
        sys.exit(1)


if __name__ == "__main__":
    try:
        test_latex_conversion()
    except Exception as e:
        print(f"Test failed with error: {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)
