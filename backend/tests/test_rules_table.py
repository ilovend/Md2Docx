"""
Table Rules Tests for Md2Docx
Run with: pytest backend/tests/test_rules_table.py -v
"""

import pytest
from docx import Document

from backend.engine.rules.table import (
    TableBorderRule,
    TableCellSpacingRule,
    TableWidthRule,
)


class TestTableBorderRule:
    """Test TableBorderRule"""

    def create_document_with_table(self):
        """Create a document with a simple table"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        return doc

    def test_apply_border_to_table(self):
        """Test applying borders to table"""
        rule = TableBorderRule()
        doc = self.create_document_with_table()

        fixes = rule.apply(doc, {})

        assert len(fixes) >= 1
        assert fixes[0]["rule_id"] == "table_border"
        assert (
            "边框" in fixes[0]["description"]
            or "border" in fixes[0]["description"].lower()
        )

    def test_border_with_custom_params(self):
        """Test border with custom parameters"""
        rule = TableBorderRule()
        doc = self.create_document_with_table()

        params = {
            "border_size": 8,
            "border_color": "000000",
        }

        fixes = rule.apply(doc, params)
        assert len(fixes) >= 1


class TestTableCellSpacingRule:
    """Test TableCellSpacingRule"""

    def create_document_with_table(self):
        """Create a document with a table"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        return doc

    def test_apply_cell_spacing(self):
        """Test applying cell spacing"""
        rule = TableCellSpacingRule()
        doc = self.create_document_with_table()

        fixes = rule.apply(doc, {})

        assert len(fixes) >= 1
        assert fixes[0]["rule_id"] == "table_cell_spacing"

    def test_cell_spacing_with_custom_params(self):
        """Test cell spacing with custom parameters"""
        rule = TableCellSpacingRule()
        doc = self.create_document_with_table()

        params = {
            "top_margin": 0.05,
            "bottom_margin": 0.05,
            "left_margin": 0.08,
            "right_margin": 0.08,
        }

        fixes = rule.apply(doc, params)
        assert len(fixes) >= 1


class TestTableWidthRule:
    """Test TableWidthRule"""

    def create_document_with_table(self):
        """Create a document with a table"""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        for i in range(2):
            for j in range(3):
                table.cell(i, j).text = f"Cell {i},{j}"
        return doc

    def test_apply_table_width(self):
        """Test applying table width"""
        rule = TableWidthRule()
        doc = self.create_document_with_table()

        fixes = rule.apply(doc, {})

        assert len(fixes) >= 1
        assert fixes[0]["rule_id"] == "table_width"

    def test_table_width_with_custom_percentage(self):
        """Test table width with custom percentage"""
        rule = TableWidthRule()
        doc = self.create_document_with_table()

        params = {
            "width_percentage": 80,
        }

        fixes = rule.apply(doc, params)
        assert len(fixes) >= 1

    def test_multiple_tables(self):
        """Test applying width to multiple tables"""
        rule = TableWidthRule()
        doc = Document()

        # Add multiple tables
        for _ in range(3):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Data"

        fixes = rule.apply(doc, {})

        # Should fix all tables
        assert len(fixes) >= 3


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
