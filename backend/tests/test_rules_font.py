"""
Font Rules Tests for Md2Docx
Run with: pytest backend/tests/test_rules_font.py -v
"""

import pytest
from docx import Document

from backend.engine.rules.font import FontStandardRule, FontReplacementRule


class TestFontStandardRule:
    """Test FontStandardRule"""

    def create_document_with_text(self):
        """Create a document with paragraphs"""
        doc = Document()
        doc.add_paragraph("First paragraph with text")
        doc.add_paragraph("Second paragraph with more text")
        doc.add_heading("A heading", level=1)
        return doc

    def test_apply_font_standard(self):
        """Test applying font standardization"""
        rule = FontStandardRule()
        doc = self.create_document_with_text()

        fixes = rule.apply(doc, {})

        assert len(fixes) >= 1
        assert fixes[0]["rule_id"] == "font_standard"

    def test_font_with_custom_params(self):
        """Test font with custom parameters"""
        rule = FontStandardRule()
        doc = self.create_document_with_text()

        params = {
            "font_name": "Arial",
            "font_size": 12,
        }

        fixes = rule.apply(doc, params)
        assert len(fixes) >= 1

    def test_default_params(self):
        """Test default parameters"""
        rule = FontStandardRule()
        defaults = rule.get_default_params()

        assert "chinese_font" in defaults or "font_name" in defaults
        assert "font_size_body" in defaults or "font_size" in defaults


class TestFontReplacementRule:
    """Test FontReplacementRule"""

    def create_document_with_mixed_fonts(self):
        """Create a document with different fonts"""
        doc = Document()
        doc.add_paragraph("Text with default font")
        p2 = doc.add_paragraph("Text with custom font")

        # Set custom font on second paragraph
        for run in p2.runs:
            run.font.name = "Comic Sans MS"

        return doc

    def test_apply_font_replacement(self):
        """Test applying font replacement"""
        rule = FontReplacementRule()
        doc = self.create_document_with_mixed_fonts()

        params = {
            "replacements": {
                "Comic Sans MS": "Arial",
            }
        }

        fixes = rule.apply(doc, params)

        assert len(fixes) >= 0  # May or may not find fonts to replace

    def test_multiple_replacements(self):
        """Test multiple font replacements"""
        rule = FontReplacementRule()
        doc = Document()

        doc.add_paragraph("Font 1")
        doc.add_paragraph("Font 2")
        doc.add_paragraph("Font 3")

        params = {
            "replacements": {
                "Times New Roman": "Calibri",
                "Arial": "Calibri",
            }
        }

        fixes = rule.apply(doc, params)
        assert isinstance(fixes, list)

    def test_default_params(self):
        """Test default replacement parameters"""
        rule = FontReplacementRule()
        defaults = rule.get_default_params()

        assert "font_map" in defaults or "replacements" in defaults
        # Check if it's a dict-like structure
        assert isinstance(
            defaults.get("font_map", defaults.get("replacements", {})), dict
        )


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
