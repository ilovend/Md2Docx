from typing import Dict, Any, List
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from backend.engine.base import BaseRule
from backend.engine.registry import registry


class ImageCenterRule(BaseRule):
    id = "image_center"
    name = "图片居中规则"
    category = "image"
    description = "将文档中包含图片的段落设置为居中对齐。"
    priority = 110

    def get_default_params(self) -> Dict[str, Any]:
        return {}

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        for i, para in enumerate(doc.paragraphs):
            has_image = False
            for run in para.runs:
                if run._element.xpath(".//a:blip"):
                    has_image = True
                    break
            if has_image:
                before_alignment = para.paragraph_format.alignment
                if before_alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    fixes.append(
                        {
                            "id": f"fix_image_center_{i}",
                            "rule_id": self.id,
                            "description": f"已将第 {i+1} 段中的图片居中",
                            "paragraph_indices": [i],
                            "before": str(before_alignment),
                            "after": str(WD_PARAGRAPH_ALIGNMENT.CENTER),
                            "location": {
                                "paragraph_index": i,
                                "type": "image_alignment",
                            },
                        }
                    )
        return fixes


class ImageResizeRule(BaseRule):
    id = "image_resize"
    name = "图片重缩放规则"
    category = "image"
    description = "自动调整过大图片的尺寸以适应页面布局。"
    priority = 120

    def get_default_params(self) -> Dict[str, Any]:
        return {"max_width": 6.0, "max_height": 8.0}

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        max_width = params.get("max_width", defaults["max_width"])
        max_height = params.get("max_height", defaults["max_height"])
        max_width_emu = int(max_width * 914400)
        max_height_emu = int(max_height * 914400)

        for i, para in enumerate(doc.paragraphs):
            for run in para.runs:
                drawings = run._element.xpath(".//w:drawing")
                for drawing in drawings:
                    for tag in ["wp:inline", "wp:anchor"]:
                        element = drawing.find(
                            f".//{tag}",
                            namespaces={
                                "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                            },
                        )
                        if element is not None:
                            extent = element.find(
                                ".//wp:extent",
                                namespaces={
                                    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                                },
                            )
                            if extent is not None:
                                cx, cy = int(extent.get("cx")), int(extent.get("cy"))
                                if cx > max_width_emu or cy > max_height_emu:
                                    before_size = {"cx": cx, "cy": cy}
                                    ratio = cx / cy
                                    if ratio > 1:
                                        new_cx = max_width_emu
                                        new_cy = int(new_cx / ratio)
                                    else:
                                        new_cy = max_height_emu
                                        new_cx = int(new_cy * ratio)
                                    extent.set("cx", str(new_cx))
                                    extent.set("cy", str(new_cy))
                                    after_size = {"cx": new_cx, "cy": new_cy}
                                    fixes.append(
                                        {
                                            "id": f"fix_image_resize_{i}",
                                            "rule_id": self.id,
                                            "description": f"已缩放第 {i+1} 段中的图片",
                                            "paragraph_indices": [i],
                                            "before": str(before_size),
                                            "after": str(after_size),
                                            "location": {
                                                "paragraph_index": i,
                                                "type": "image_extent",
                                                "max_width_in": max_width,
                                                "max_height_in": max_height,
                                            },
                                        }
                                    )
        return fixes


class ImageCaptionRule(BaseRule):
    id = "image_caption"
    name = "图片题注规则"
    category = "image"
    description = "为图片添加默认的题注占位符或序号。"
    priority = 130

    def get_default_params(self) -> Dict[str, Any]:
        return {"caption_format": "图 {number}", "caption_font_size": 10}

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        image_counter = 1

        for i, para in enumerate(doc.paragraphs):
            # Check if paragraph contains images
            has_image = False
            for run in para.runs:
                if run._element.xpath(".//a:blip"):
                    has_image = True
                    break

            if has_image:
                # Add caption paragraph after the image paragraph
                caption_format = params.get(
                    "caption_format", defaults["caption_format"]
                )
                caption_text = caption_format.format(number=image_counter)

                # Create and insert caption paragraph
                caption_para = doc.add_paragraph(caption_text)
                caption_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Set caption font size
                font_size = params.get(
                    "caption_font_size", defaults["caption_font_size"]
                )
                for run in caption_para.runs:
                    run.font.size = Pt(font_size)

                image_counter += 1
                fixes.append(
                    {
                        "id": f"fix_image_caption_{i}",
                        "rule_id": self.id,
                        "description": f"已为第 {i+1} 段中的图片添加题注 '{caption_text}'",
                        "paragraph_indices": [i],
                        "before": None,
                        "after": caption_text,
                        "location": {
                            "paragraph_index": i,
                            "type": "image_caption",
                        },
                    }
                )

        return fixes


registry.register(ImageCenterRule())
registry.register(ImageResizeRule())
registry.register(ImageCaptionRule())
