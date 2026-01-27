from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from backend.engine.base import BaseRule
from backend.engine.registry import registry

class FontStandardRule(BaseRule):
    id = "font_standard"
    name = "标准字体规则"
    category = "font"
    description = "为文档中的所有段落设置标准的中西文字体和大小。"
    priority = 10

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "western_font": "Arial",
            "chinese_font": "SimSun",
            "font_size_body": 12
        }

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        western_font = params.get("western_font", defaults["western_font"])
        chinese_font = params.get("chinese_font", defaults["chinese_font"])
        font_size = params.get("font_size_body", defaults["font_size_body"])

        for i, para in enumerate(doc.paragraphs):
            changed = False
            for run in para.runs:
                if run.font.name != western_font:
                    run.font.name = western_font
                    changed = True
                if run._element.rPr is not None:
                    rFonts = run._element.rPr.rFonts
                    if rFonts is not None:
                        rFonts.set(qn("w:eastAsia"), chinese_font)
                if run.font.size != Pt(font_size):
                    run.font.size = Pt(font_size)
                    changed = True

            if changed:
                fixes.append({
                    "id": f"fix_font_{i}",
                    "rule_id": self.id,
                    "description": f"已应用字体 {western_font}/{chinese_font} 大小 {font_size}pt",
                    "paragraph_indices": [i],
                })
        return fixes

class FontColorRule(BaseRule):
    id = "font_color"
    name = "字体颜色规则"
    category = "font"
    description = "将文档全文的字体颜色设置为指定值。"
    priority = 70

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "text_color": "000000"
        }

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        text_color = params.get("text_color", defaults["text_color"])

        try:
            r = int(text_color[0:2], 16)
            g = int(text_color[2:4], 16)
            b = int(text_color[4:6], 16)
            color = RGBColor(r, g, b)
        except Exception:
            color = RGBColor(0, 0, 0)

        affected_indices = []
        for i, para in enumerate(doc.paragraphs):
            changed = False
            for run in para.runs:
                if run.font.color.rgb != color:
                    run.font.color.rgb = color
                    changed = True
            if changed:
                affected_indices.append(i)

        if affected_indices:
            fixes.append({
                "id": "fix_font_color_all",
                "rule_id": self.id,
                "description": f"已应用字体颜色 #{text_color}",
                "paragraph_indices": affected_indices,
            })
        return fixes

class FontReplacementRule(BaseRule):
    id = "font_replacement"
    name = "字体替换规则"
    category = "font"
    description = "将非标准字体替换为指定的替代字体。"
    priority = 80

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "font_map": {
                "Microsoft YaHei": "SimSun",
                "微软雅黑": "SimSun",
                "宋体": "SimSun",
                "黑体": "SimHei",
                "楷体": "KaiTi",
            },
            "default_western_font": "Arial",
            "default_chinese_font": "SimSun"
        }

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        font_map = params.get("font_map", defaults["font_map"])
        default_western_font = params.get("default_western_font", defaults["default_western_font"])
        default_chinese_font = params.get("default_chinese_font", defaults["default_chinese_font"])

        affected_indices = []
        for i, para in enumerate(doc.paragraphs):
            changed = False
            for run in para.runs:
                if run.font.name:
                    current_font = run.font.name
                    if current_font in font_map:
                        new_font = font_map[current_font]
                        if run.font.name != new_font:
                            run.font.name = new_font
                            changed = True
                    elif run.font.name != default_western_font:
                        run.font.name = default_western_font
                        changed = True
                else:
                    run.font.name = default_western_font
                    changed = True

                if run._element.rPr is not None:
                    rFonts = run._element.rPr.rFonts
                    if rFonts is not None:
                        east_asia_font = rFonts.get(qn("w:eastAsia"))
                        if east_asia_font:
                            if east_asia_font in font_map:
                                new_font = font_map[east_asia_font]
                                if rFonts.get(qn("w:eastAsia")) != new_font:
                                    rFonts.set(qn("w:eastAsia"), new_font)
                                    changed = True
                            elif rFonts.get(qn("w:eastAsia")) != default_chinese_font:
                                rFonts.set(qn("w:eastAsia"), default_chinese_font)
                                changed = True
                        else:
                            rFonts.set(qn("w:eastAsia"), default_chinese_font)
                            changed = True

            if changed:
                affected_indices.append(i)

        if affected_indices:
            fixes.append({
                "id": "fix_font_replacement_all",
                "rule_id": self.id,
                "description": f"已将非标准字体替换为 {default_western_font}/{default_chinese_font}",
                "paragraph_indices": affected_indices,
            })
        return fixes

# 自动注册
registry.register(FontStandardRule())
registry.register(FontColorRule())
registry.register(FontReplacementRule())
