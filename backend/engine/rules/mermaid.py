from typing import Dict, Any, List
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from backend.engine.base import BaseRule
from backend.engine.registry import registry
import re
import subprocess
import tempfile
import os


class MermaidRenderRule(BaseRule):
    id = "mermaid_render"
    name = "Mermaid图表渲染规则"
    category = "image"
    description = "将Markdown中的Mermaid代码块转换为图片并插入文档。"
    priority = 135

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "output_format": "png",
            "scale": 2,
            "theme": "default",
            "background_color": "white",
        }

    def _detect_mermaid_code(self, text: str) -> bool:
        """检测段落是否包含Mermaid代码块"""
        # 检测 ```mermaid 或 ~~~mermaid 代码块
        mermaid_pattern = r"```mermaid|~~~mermaid"
        return bool(re.search(mermaid_pattern, text, re.IGNORECASE))

    def _extract_mermaid_code(self, text: str) -> str:
        """提取Mermaid代码内容"""
        # 提取代码块内容
        pattern = r"```mermaid\s*(.*?)\s*```|~~~mermaid\s*(.*?)\s*~~~"
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            return match.group(1) or match.group(2)
        return ""

    def _render_mermaid_to_image(
        self, mermaid_code: str, params: Dict[str, Any]
    ) -> str:
        """
        使用mermaid-cli渲染Mermaid代码为图片
        需要安装: npm install -g @mermaid-js/mermaid-cli
        """
        output_format = params.get("output_format", "png")
        scale = params.get("scale", 2)
        theme = params.get("theme", "default")
        bg_color = params.get("background_color", "white")

        # 创建临时文件
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".mmd", delete=False, encoding="utf-8"
        ) as mmd_file:
            mmd_file.write(mermaid_code)
            mmd_path = mmd_file.name

        output_path = mmd_path.replace(".mmd", f".{output_format}")

        try:
            # 调用mermaid-cli (mmdc命令)
            cmd = [
                "mmdc",
                "-i",
                mmd_path,
                "-o",
                output_path,
                "-t",
                theme,
                "-b",
                bg_color,
                "-s",
                str(scale),
            ]

            result = subprocess.run(
                cmd, capture_output=True, text=True, timeout=30, check=True
            )

            if os.path.exists(output_path):
                return output_path
            else:
                raise RuntimeError(f"Mermaid rendering failed: {result.stderr}")

        except subprocess.CalledProcessError as e:
            raise RuntimeError(
                f"Mermaid CLI error: {e.stderr}. "
                "Please ensure mermaid-cli is installed: npm install -g @mermaid-js/mermaid-cli"
            )
        except FileNotFoundError:
            raise RuntimeError(
                "mermaid-cli (mmdc) not found. "
                "Please install: npm install -g @mermaid-js/mermaid-cli"
            )
        finally:
            # 清理临时mmd文件
            if os.path.exists(mmd_path):
                os.remove(mmd_path)

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text

            if self._detect_mermaid_code(text):
                try:
                    # 提取Mermaid代码
                    mermaid_code = self._extract_mermaid_code(text)
                    if not mermaid_code.strip():
                        continue

                    # 渲染为图片
                    image_path = self._render_mermaid_to_image(mermaid_code, params)

                    # 在当前段落位置插入图片
                    # 清空当前段落内容
                    para.clear()

                    # 添加图片
                    run = para.add_run()
                    run.add_picture(image_path, width=Inches(6.0))

                    # 设置居中对齐
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # 清理临时图片文件
                    if os.path.exists(image_path):
                        os.remove(image_path)

                    fixes.append(
                        {
                            "id": f"fix_mermaid_render_{i}",
                            "rule_id": self.id,
                            "description": f"已将第 {i+1} 段中的Mermaid图表渲染为图片",
                            "paragraph_indices": [i],
                            "before": text[:100] + "..." if len(text) > 100 else text,
                            "after": "[Mermaid图表]",
                            "location": {
                                "paragraph_index": i,
                                "type": "mermaid_diagram",
                            },
                        }
                    )

                except Exception as e:
                    # 如果渲染失败，记录错误但继续处理
                    fixes.append(
                        {
                            "id": f"fix_mermaid_render_error_{i}",
                            "rule_id": self.id,
                            "description": f"第 {i+1} 段Mermaid渲染失败: {str(e)}",
                            "paragraph_indices": [i],
                            "before": text[:100] + "..." if len(text) > 100 else text,
                            "after": None,
                            "location": {
                                "paragraph_index": i,
                                "type": "mermaid_error",
                                "error": str(e),
                            },
                        }
                    )

        return fixes


# 注册规则
registry.register(MermaidRenderRule())
