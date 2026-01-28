from typing import Dict, Any, List
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from backend.engine.base import BaseRule
from backend.engine.registry import registry


class TableBorderRule(BaseRule):
    id = "table_border"
    name = "表格边框规则"
    category = "table"
    description = "为所有表格应用统一的边框样式和表头背景色。"
    priority = 20

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "border_size": 4,
            "border_color": "000000",
            "add_table_header_format": True,
            "table_header_bg_color": "E3E3E3",
        }

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        border_size = params.get("border_size", defaults["border_size"])
        border_color = params.get("border_color", defaults["border_color"])

        for i, table in enumerate(doc.tables):
            self._set_table_borders(table, border_size, border_color)
            desc = f"已为表格 {i+1} 应用 {border_size}pt 边框"

            if (
                params.get(
                    "add_table_header_format", defaults["add_table_header_format"]
                )
                and len(table.rows) > 0
            ):
                header_bg = params.get(
                    "table_header_bg_color", defaults["table_header_bg_color"]
                )
                self._set_row_shading(table.rows[0], header_bg)
                desc += f" 且设置表头背景色 #{header_bg}"

            fixes.append(
                {
                    "id": f"fix_table_{i}",
                    "rule_id": self.id,
                    "description": desc,
                    "table_indices": [i],
                    "before": None,
                    "after": desc,
                    "location": {
                        "table_index": i,
                        "type": "table_borders",
                        "border_size": border_size,
                        "border_color": border_color,
                    },
                }
            )
        return fixes

    def _set_table_borders(self, table, size, color):
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(size))
            border.set(qn("w:color"), color)
            tblBorders.append(border)
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _set_row_shading(self, row, color):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), color)
            tcPr.append(shd)


class TableWidthRule(BaseRule):
    id = "table_width"
    name = "表格宽度规则"
    category = "table"
    description = "将表格宽度设置为指定页面宽度百分比，并可选启用自动列宽。"
    priority = 35

    def get_default_params(self) -> Dict[str, Any]:
        return {"table_width_percent": 95, "auto_adjust_columns": True}

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes: List[Dict[str, Any]] = []
        defaults = self.get_default_params()
        width_percent = int(
            params.get("table_width_percent", defaults["table_width_percent"])
        )
        auto_adjust = bool(
            params.get("auto_adjust_columns", defaults["auto_adjust_columns"])
        )

        # Approximate usable page width from first section
        usable_width_in = None
        if doc.sections:
            sec = doc.sections[0]
            try:
                usable_width = sec.page_width - sec.left_margin - sec.right_margin
                usable_width_in = usable_width.inches
            except Exception:
                usable_width_in = None

        for i, table in enumerate(doc.tables):
            before = {
                "table_width_percent": None,
                "usable_width_in": usable_width_in,
            }

            target_width_in = None
            if usable_width_in is not None:
                target_width_in = usable_width_in * (width_percent / 100.0)

            # Set table width via tblW if possible
            if target_width_in is not None:
                tbl = table._tbl
                tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
                tblW = OxmlElement("w:tblW")
                # dxa: 1 inch = 1440 twips
                tblW.set(qn("w:type"), "dxa")
                tblW.set(qn("w:w"), str(int(target_width_in * 1440)))
                tblPr.append(tblW)
                if tbl.tblPr is None:
                    tbl.insert(0, tblPr)

            if auto_adjust:
                # reuse the existing logic: set autofit layout
                tbl = table._tbl
                tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
                tblLayout = OxmlElement("w:tblLayout")
                tblLayout.set(qn("w:type"), "autofit")
                tblPr.append(tblLayout)
                if tbl.tblPr is None:
                    tbl.insert(0, tblPr)

            desc = f"已将表格 {i+1} 宽度设为页面的 {width_percent}%"
            fixes.append(
                {
                    "id": f"fix_table_width_{i}",
                    "rule_id": self.id,
                    "description": desc,
                    "table_indices": [i],
                    "before": str(before),
                    "after": str(
                        {
                            "table_width_percent": width_percent,
                            "auto_adjust_columns": auto_adjust,
                            "target_width_in": target_width_in,
                        }
                    ),
                    "location": {
                        "table_index": i,
                        "type": "table_width",
                        "width_percent": width_percent,
                    },
                }
            )

        return fixes


class TableCellSpacingRule(BaseRule):
    id = "table_cell_spacing"
    name = "表格单元格间距规则"
    category = "table"
    description = "为所有表格设置统一的单元格内边距。"
    priority = 30

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "cell_margin_top": 50,
            "cell_margin_left": 50,
            "cell_margin_bottom": 50,
            "cell_margin_right": 50,
        }

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        tm = params.get("cell_margin_top", defaults["cell_margin_top"])
        lm = params.get("cell_margin_left", defaults["cell_margin_left"])
        bm = params.get("cell_margin_bottom", defaults["cell_margin_bottom"])
        rm = params.get("cell_margin_right", defaults["cell_margin_right"])

        for i, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcMar = OxmlElement("w:tcMar")
                    for side, val in zip(
                        ["top", "left", "bottom", "right"], [tm, lm, bm, rm]
                    ):
                        node = OxmlElement(f"w:{side}")
                        node.set(qn("w:w"), str(val))
                        node.set(qn("w:type"), "dxa")
                        tcMar.append(node)
                    tcPr.append(tcMar)
            fixes.append(
                {
                    "id": f"fix_table_cell_spacing_{i}",
                    "rule_id": self.id,
                    "description": f"已为表格 {i+1} 应用单元格间距",
                    "table_indices": [i],
                    "before": None,
                    "after": str({"top": tm, "left": lm, "bottom": bm, "right": rm}),
                    "location": {
                        "table_index": i,
                        "type": "table_cell_margins",
                        "top": tm,
                        "left": lm,
                        "bottom": bm,
                        "right": rm,
                    },
                }
            )
        return fixes


class TableColumnWidthRule(BaseRule):
    id = "table_column_width"
    name = "表格列宽自适应规则"
    category = "table"
    description = "将表格设置为自动调整列宽以适应内容。"
    priority = 40

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        for i, table in enumerate(doc.tables):
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
            tblLayout = OxmlElement("w:tblLayout")
            tblLayout.set(qn("w:type"), "autofit")
            tblPr.append(tblLayout)
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    for element in tcPr.findall(qn("w:tcW")):
                        tcPr.remove(element)
            fixes.append(
                {
                    "id": f"fix_table_column_width_{i}",
                    "rule_id": self.id,
                    "description": f"已为表格 {i+1} 应用自动列宽",
                    "table_indices": [i],
                    "before": None,
                    "after": "autofit",
                    "location": {
                        "table_index": i,
                        "type": "table_layout",
                        "layout": "autofit",
                    },
                }
            )
        return fixes


registry.register(TableBorderRule())
registry.register(TableCellSpacingRule())
registry.register(TableWidthRule())
registry.register(TableColumnWidthRule())
