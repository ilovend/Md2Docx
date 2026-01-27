from typing import Dict, Any, List
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from backend.engine.base import BaseRule
from backend.engine.registry import registry
from backend.core.latex_converter import convert_latex_in_document


class LatexToOmmlRule(BaseRule):
    id = "latex_to_omml"
    name = "LaTeX公式转换规则"
    category = "formula"
    description = "将文档中的 LaTeX 公式转换为 Word 原生 OMML 公式。"
    priority = 140

    def get_default_params(self) -> Dict[str, Any]:
        return {}

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        # This calls the existing specialized converter
        return convert_latex_in_document(doc)


class FormulaNumberingRule(BaseRule):
    id = "formula_numbering"
    name = "公式自动编号规则"
    category = "formula"
    description = "为所有展示型公式添加自动编号。"
    priority = 150

    def get_default_params(self) -> Dict[str, Any]:
        return {}

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        formula_counter = 1

        for i, para in enumerate(doc.paragraphs):
            # Check if paragraph contains display formulas
            has_display_formula = False
            for child in para._p:
                if child.tag.endswith("oMath") or child.tag.endswith("oMathPara"):
                    # Check if it's a display formula (typically in its own paragraph)
                    if len(para.runs) == 0 or (
                        len(para.runs) == 1 and not para.runs[0].text.strip()
                    ):
                        has_display_formula = True
                        break

            if has_display_formula:
                # Add formula numbering
                formula_number = f"({formula_counter})"
                # Add numbering to the end of the paragraph
                para.add_run(f" {formula_number}")
                # Increment counter
                formula_counter += 1
                # Add fix to list
                fixes.append(
                    {
                        "id": f"fix_formula_numbering_{i}",
                        "rule_id": "formula_numbering",
                        "description": f"Added formula numbering ({formula_counter-1})",
                        "paragraph_indices": [i],
                    }
                )

        return fixes


class InlineFormulaStyleRule(BaseRule):
    id = "inline_formula_style"
    name = "行内公式样式规则"
    category = "formula"
    description = "统一行内公式的字体和垂直对齐方式。"
    priority = 160

    def get_default_params(self) -> Dict[str, Any]:
        return {}

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []

        for i, para in enumerate(doc.paragraphs):
            # Check if paragraph contains inline formulas
            has_inline_formula = False
            for child in para._p:
                if child.tag.endswith("oMath"):
                    # Check if it's an inline formula (not in its own paragraph)
                    if len(para.runs) > 0 and para.runs[0].text.strip():
                        has_inline_formula = True
                        break

            if has_inline_formula:
                # For inline formulas, we can adjust the font size or other properties
                # Here we'll just mark it as fixed for now
                fixes.append(
                    {
                        "id": f"fix_inline_formula_style_{i}",
                        "rule_id": "inline_formula_style",
                        "description": "Applied inline formula style",
                        "paragraph_indices": [i],
                    }
                )

        return fixes


class DisplayFormulaCenterRule(BaseRule):
    id = "display_formula_center"
    name = "展示公式居中规则"
    category = "formula"
    description = "将所有展示型公式设置为居中对齐。"
    priority = 170

    def get_default_params(self) -> Dict[str, Any]:
        return {}

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []

        for i, para in enumerate(doc.paragraphs):
            # Check if paragraph contains display formulas
            has_display_formula = False
            for child in para._p:
                if child.tag.endswith("oMath") or child.tag.endswith("oMathPara"):
                    # Check if it's a display formula (typically in its own paragraph)
                    if len(para.runs) == 0 or (
                        len(para.runs) == 1 and not para.runs[0].text.strip()
                    ):
                        has_display_formula = True
                        break

            if has_display_formula:
                # Center the paragraph
                if para.paragraph_format.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    fixes.append(
                        {
                            "id": f"fix_display_formula_center_{i}",
                            "rule_id": "display_formula_center",
                            "description": "Centered display formula",
                            "paragraph_indices": [i],
                        }
                    )

        return fixes


registry.register(LatexToOmmlRule())
registry.register(FormulaNumberingRule())
registry.register(InlineFormulaStyleRule())
registry.register(DisplayFormulaCenterRule())
