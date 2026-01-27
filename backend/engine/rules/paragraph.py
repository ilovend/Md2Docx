from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt
from backend.engine.base import BaseRule
from backend.engine.registry import registry

class ParagraphSpacingRule(BaseRule):
    id = "paragraph_spacing"
    name = "段落间距规则"
    category = "paragraph"
    description = "设置文档段落的行间距和段前段后距离。"
    priority = 50

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "line_spacing": 1.5,
            "space_before": 0,
            "space_after": 6
        }

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        line_spacing = params.get("line_spacing", defaults["line_spacing"])
        space_before = params.get("space_before", defaults["space_before"])
        space_after = params.get("space_after", defaults["space_after"])

        affected_indices = []
        for i, para in enumerate(doc.paragraphs):
            pf = para.paragraph_format
            pf.line_spacing = line_spacing
            pf.space_before = Pt(space_before)
            pf.space_after = Pt(space_after)
            affected_indices.append(i)

        if affected_indices:
            fixes.append({
                "id": "fix_spacing_all",
                "rule_id": self.id,
                "description": f"已应用行间距 {line_spacing}x，段前 {space_before}pt，段后 {space_after}pt",
                "paragraph_indices": affected_indices,
            })
        return fixes

class FirstLineIndentRule(BaseRule):
    id = "first_line_indent"
    name = "首行缩进规则"
    category = "paragraph"
    description = "为正文段落应用首行缩进。"
    priority = 90

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "indent_size": 2
        }

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        indent_size = params.get("indent_size", defaults["indent_size"])
        indent_pt = Pt(indent_size * 12)
        heading_styles = ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"]

        affected_indices = []
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name in heading_styles:
                continue
            if not para.text.strip():
                continue
            pf = para.paragraph_format
            if pf.first_line_indent != indent_pt:
                pf.first_line_indent = indent_pt
                affected_indices.append(i)

        if affected_indices:
            fixes.append({
                "id": "fix_first_line_indent",
                "rule_id": self.id,
                "description": f"已应用 {indent_size} 字符首行缩进",
                "paragraph_indices": affected_indices,
            })
        return fixes

class TitleBoldRule(BaseRule):
    id = "title_bold"
    name = "标题加粗规则"
    category = "heading"
    description = "确保所有标题层级都应用加粗样式。"
    priority = 60

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        heading_styles = ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"]
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name in heading_styles:
                changed = False
                for run in para.runs:
                    if not run.bold:
                        run.bold = True
                        changed = True
                if changed:
                    fixes.append({
                        "id": f"fix_title_bold_{i}",
                        "rule_id": self.id,
                        "description": f"已将标题 '{para.style.name}' 加粗",
                        "paragraph_indices": [i],
                    })
        return fixes

class HeadingStyleRule(BaseRule):
    id = "heading_style"
    name = "标题样式规则"
    category = "heading"
    description = "统一各级标题的字体大小和样式。"
    priority = 100

    def get_default_params(self) -> Dict[str, Any]:
        return {
            "h1_size": 22,
            "h2_size": 16,
            "h3_size": 14
        }

    def apply(self, doc: Document, params: Dict[str, Any]) -> List[Dict[str, Any]]:
        fixes = []
        defaults = self.get_default_params()
        h1_size = params.get("h1_size", defaults["h1_size"])
        h2_size = params.get("h2_size", defaults["h2_size"])
        h3_size = params.get("h3_size", defaults["h3_size"])

        size_map = {
            "Heading 1": h1_size,
            "Heading 2": h2_size,
            "Heading 3": h3_size,
            "Title": h1_size + 4,
        }
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name in size_map:
                target_size = size_map[para.style.name]
                changed = False
                for run in para.runs:
                    if run.font.size != Pt(target_size):
                        run.font.size = Pt(target_size)
                        run.bold = True
                        changed = True
                if changed:
                    fixes.append({
                        "id": f"fix_heading_{i}",
                        "rule_id": self.id,
                        "description": f"已应用 {target_size}pt 到标题 {para.style.name}",
                        "paragraph_indices": [i],
                    })
        return fixes

registry.register(ParagraphSpacingRule())
registry.register(FirstLineIndentRule())
registry.register(TitleBoldRule())
registry.register(HeadingStyleRule())
