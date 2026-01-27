"""
Markdown to Word Converter
Converts Markdown files to Word documents with proper formatting.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownConverter:
    """Converts Markdown content to Word document."""

    def __init__(self):
        self.doc = None
        self.current_list_level = 0

    def convert(self, markdown_path: Path, output_path: Path) -> dict:
        """
        Convert Markdown file to Word document.

        Args:
            markdown_path: Path to the Markdown file
            output_path: Path for the output Word document

        Returns:
            dict with conversion statistics
        """
        with open(markdown_path, "r", encoding="utf-8") as f:
            content = f.read()

        self.doc = Document()
        stats = {
            "headings": 0,
            "paragraphs": 0,
            "code_blocks": 0,
            "lists": 0,
            "tables": 0,
            "blockquotes": 0,
        }

        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Code block
            if line.strip().startswith("```"):
                code_lines = []
                lang = line.strip()[3:]
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block("\n".join(code_lines), lang)
                stats["code_blocks"] += 1
                i += 1
                continue

            # Heading
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self._add_heading(text, level)
                stats["headings"] += 1
                i += 1
                continue

            # Blockquote
            if line.strip().startswith(">"):
                quote_text = line.strip()[1:].strip()
                self._add_blockquote(quote_text)
                stats["blockquotes"] += 1
                i += 1
                continue

            # Unordered list
            if re.match(r"^[\s]*[-*+]\s+", line):
                self._add_list_item(line, ordered=False)
                stats["lists"] += 1
                i += 1
                continue

            # Ordered list
            if re.match(r"^[\s]*\d+\.\s+", line):
                self._add_list_item(line, ordered=True)
                stats["lists"] += 1
                i += 1
                continue

            # Table
            if "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
                table_lines = [line]
                i += 1
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(table_lines)
                stats["tables"] += 1
                continue

            # Regular paragraph
            self._add_paragraph(line)
            stats["paragraphs"] += 1
            i += 1

        self.doc.save(output_path)
        return stats

    def _add_heading(self, text: str, level: int):
        """Add heading to document."""
        style_name = f"Heading {level}"
        para = self.doc.add_paragraph(style=style_name)
        self._add_formatted_text(para, text)

    def _add_paragraph(self, text: str):
        """Add paragraph with inline formatting."""
        para = self.doc.add_paragraph()
        self._add_formatted_text(para, text)

    def _add_formatted_text(self, para, text: str):
        """Parse and add formatted text (bold, italic, code, links)."""
        # Pattern for inline formatting
        patterns = [
            (r"\*\*\*(.+?)\*\*\*", "bold_italic"),  # Bold italic
            (r"\*\*(.+?)\*\*", "bold"),  # Bold
            (r"\*(.+?)\*", "italic"),  # Italic
            (r"`(.+?)`", "code"),  # Inline code
            (r"\[(.+?)\]\((.+?)\)", "link"),  # Link
        ]

        # Simple approach: just add text with basic formatting detection
        remaining = text

        while remaining:
            # Find the earliest match
            earliest_match = None
            earliest_type = None
            earliest_pos = len(remaining)

            for pattern, fmt_type in patterns:
                match = re.search(pattern, remaining)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_type = fmt_type
                    earliest_pos = match.start()

            if earliest_match:
                # Add text before match
                if earliest_pos > 0:
                    para.add_run(remaining[:earliest_pos])

                # Add formatted text
                if earliest_type == "bold":
                    run = para.add_run(earliest_match.group(1))
                    run.bold = True
                elif earliest_type == "italic":
                    run = para.add_run(earliest_match.group(1))
                    run.italic = True
                elif earliest_type == "bold_italic":
                    run = para.add_run(earliest_match.group(1))
                    run.bold = True
                    run.italic = True
                elif earliest_type == "code":
                    run = para.add_run(earliest_match.group(1))
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                elif earliest_type == "link":
                    run = para.add_run(earliest_match.group(1))
                    run.font.color.rgb = RGBColor(0, 102, 204)
                    run.underline = True

                remaining = remaining[earliest_match.end() :]
            else:
                # No more matches, add remaining text
                para.add_run(remaining)
                break

    def _add_code_block(self, code: str, language: str = ""):
        """Add code block with monospace font."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)

        # Add background shading
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F5F5F5")
        para._element.get_or_add_pPr().append(shading)

        run = para.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def _add_blockquote(self, text: str):
        """Add blockquote with left border styling."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)

        run = para.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)

    def _add_list_item(self, line: str, ordered: bool = False):
        """Add list item."""
        # Extract text from list item
        if ordered:
            match = re.match(r"^[\s]*\d+\.\s+(.+)$", line)
        else:
            match = re.match(r"^[\s]*[-*+]\s+(.+)$", line)

        if match:
            text = match.group(1)
            style = "List Number" if ordered else "List Bullet"
            para = self.doc.add_paragraph(style=style)
            self._add_formatted_text(para, text)

    def _add_table(self, lines: list):
        """Add table from Markdown table syntax."""
        # Parse header row
        headers = [cell.strip() for cell in lines[0].split("|") if cell.strip()]

        # Skip separator row (index 1)
        # Parse data rows
        data_rows = []
        for line in lines[2:]:
            if "|" in line:
                cells = [cell.strip() for cell in line.split("|") if cell.strip()]
                if cells:
                    data_rows.append(cells)

        if not headers:
            return

        # Create table
        table = self.doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
        table.style = "Table Grid"

        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            if i < len(header_row.cells):
                cell = header_row.cells[i]
                cell.text = header
                # Bold header
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

        # Add data
        for row_idx, row_data in enumerate(data_rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = cell_text


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> dict:
    """
    Convenience function to convert Markdown to Word.

    Args:
        md_path: Path to Markdown file
        docx_path: Path for output Word file

    Returns:
        Conversion statistics
    """
    converter = MarkdownConverter()
    return converter.convert(md_path, docx_path)
