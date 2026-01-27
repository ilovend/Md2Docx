"""
LaTeX to OMML Converter
Converts LaTeX math expressions to Word OMML format.
"""

import re
from pathlib import Path
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Try to import latex2mathml, fall back to basic conversion if not available
try:
    import latex2mathml.converter as latex2mathml

    HAS_LATEX2MATHML = True
except ImportError:
    HAS_LATEX2MATHML = False

# OMML namespace
OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# MathML to OMML XSLT (simplified version)
# In production, use the full Microsoft MML2OMML.xsl stylesheet
MATHML_TO_OMML_XSLT = """<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0"
    xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:mml="http://www.w3.org/1998/Math/MathML">
    
    <xsl:output method="xml" indent="yes"/>
    
    <xsl:template match="mml:math">
        <m:oMath>
            <xsl:apply-templates/>
        </m:oMath>
    </xsl:template>
    
    <xsl:template match="mml:mrow">
        <xsl:apply-templates/>
    </xsl:template>
    
    <xsl:template match="mml:mi|mml:mn|mml:mo|mml:mtext">
        <m:r>
            <m:t><xsl:value-of select="."/></m:t>
        </m:r>
    </xsl:template>
    
    <xsl:template match="mml:mfrac">
        <m:f>
            <m:num><xsl:apply-templates select="*[1]"/></m:num>
            <m:den><xsl:apply-templates select="*[2]"/></m:den>
        </m:f>
    </xsl:template>
    
    <xsl:template match="mml:msup">
        <m:sSup>
            <m:e><xsl:apply-templates select="*[1]"/></m:e>
            <m:sup><xsl:apply-templates select="*[2]"/></m:sup>
        </m:sSup>
    </xsl:template>
    
    <xsl:template match="mml:msub">
        <m:sSub>
            <m:e><xsl:apply-templates select="*[1]"/></m:e>
            <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
        </m:sSub>
    </xsl:template>
    
    <xsl:template match="mml:msqrt">
        <m:rad>
            <m:radPr><m:degHide m:val="1"/></m:radPr>
            <m:deg/>
            <m:e><xsl:apply-templates/></m:e>
        </m:rad>
    </xsl:template>
    
    <xsl:template match="mml:munderover|mml:msubsup">
        <m:sSubSup>
            <m:e><xsl:apply-templates select="*[1]"/></m:e>
            <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
            <m:sup><xsl:apply-templates select="*[3]"/></m:sup>
        </m:sSubSup>
    </xsl:template>
    
    <xsl:template match="*">
        <xsl:apply-templates/>
    </xsl:template>
</xsl:stylesheet>
"""


class LaTeXConverter:
    """Converts LaTeX math expressions to Word OMML format."""

    def __init__(self):
        self.xslt_transform = None
        try:
            xslt_doc = etree.fromstring(MATHML_TO_OMML_XSLT.encode())
            self.xslt_transform = etree.XSLT(xslt_doc)
        except Exception as e:
            print(f"Warning: Could not initialize XSLT transform: {e}")

    def latex_to_mathml(self, latex: str) -> str:
        """Convert LaTeX to MathML."""
        if HAS_LATEX2MATHML:
            try:
                mathml = latex2mathml.convert(latex)
                return mathml
            except Exception as e:
                print(f"latex2mathml error: {e}")
                return None
        return None

    def mathml_to_omml(self, mathml: str) -> etree.Element:
        """Convert MathML to OMML using XSLT."""
        if not self.xslt_transform:
            return None

        try:
            # Parse MathML
            mathml_doc = etree.fromstring(mathml.encode())

            # Transform to OMML
            omml_doc = self.xslt_transform(mathml_doc)

            return omml_doc.getroot()
        except Exception as e:
            print(f"MathML to OMML error: {e}")
            return None

    def latex_to_omml(self, latex: str) -> etree.Element:
        """Convert LaTeX directly to OMML."""
        mathml = self.latex_to_mathml(latex)
        if mathml:
            return self.mathml_to_omml(mathml)
        return None

    def create_simple_omml(self, latex: str) -> etree.Element:
        """Create a simple OMML element for basic math (fallback)."""
        # Create basic OMML structure
        omath = OxmlElement("m:oMath")
        omath.set(qn("xmlns:m"), OMML_NS)

        # Add the LaTeX as text (will display but not be editable as equation)
        r = OxmlElement("m:r")
        t = OxmlElement("m:t")
        t.text = latex
        r.append(t)
        omath.append(r)

        return omath

    def replace_latex_with_omml(self, para, match_start, match_end, omml_element):
        """
        Replace text in paragraph with OMML element.
        This is a complex operation in python-docx because we need to split runs.
        """
        p = para._p

        # 1. Remove the original text
        # This is a simplification. robust implementation needs to handle runs carefully.
        # For now, we'll assume the simple case where we can just clear the paragraph text
        # around the match, but that's risky.

        # Better approach:
        # 1. Clear the paragraph text (or the specific runs)
        # 2. Rebuild the paragraph with: Text Before -> OMML -> Text After

        original_text = para.text
        text_before = original_text[:match_start]
        text_after = original_text[match_end:]

        # Clear all runs
        for run in para.runs:
            p.remove(run._r)

        # Add text before
        if text_before:
            para.add_run(text_before)

        # Add OMML
        # We need to wrap OMML in a run-like structure or append directly to p
        # Microsoft Word expects m:oMathPara for display math or m:oMath for inline
        # inside the paragraph.

        # If display math, we might want a new paragraph, but here we stay in same para
        # to minimize disruption.

        p.append(omml_element)

        # Add text after
        if text_after:
            para.add_run(text_after)

    def process_document(self, doc: Document) -> list:
        """
        Find and convert LaTeX expressions in document.

        Returns list of fixes applied.
        """
        fixes = []

        # LaTeX patterns
        # Order matters! Check for display math first
        patterns = [
            (r"\$\$(.+?)\$\$", "display"),  # Display math $$...$$
            (r"\\\[(.+?)\\\]", "display"),  # Display math \[...\]
            (r"\\begin\{equation\}(.+?)\\end\{equation\}", "display"),  # equation env
            (r"\$(.+?)\$", "inline"),  # Inline math $...$
        ]

        # Iterate backwards to avoid index shifting issues when replacing
        # But python-docx doesn't support easy replacement, so we might need a different strategy.
        # Since we are modifying the paragraph content entirely in replace_latex_with_omml,
        # we need to be careful not to invalidate iteration.

        # Safe strategy: Find one match, replace it, restart search in that paragraph?
        # Or just one match per paragraph for MVP safety.

        for para_idx, para in enumerate(doc.paragraphs):
            # Skip if empty
            if not para.text.strip():
                continue

            original_text = para.text

            for pattern, math_type in patterns:
                # Find *one* match per paragraph to keep it safe for MVP
                # A full recurring search is harder because we modify valid ranges
                match = re.search(pattern, original_text, re.DOTALL)

                if match:
                    latex_expr = match.group(1).strip()
                    full_match = match.group(0)
                    start, end = match.span()

                    # Try OMML conversion
                    omml = self.latex_to_omml(latex_expr)

                    if omml is not None:
                        try:
                            self.replace_latex_with_omml(para, start, end, omml)

                            fixes.append(
                                {
                                    "id": f"fix_latex_{para_idx}",
                                    "rule_id": "latex_to_omml",
                                    "description": f"Converted LaTeX: {latex_expr[:30]}",
                                    "paragraph_indices": [para_idx],
                                }
                            )

                            # Break inner loop to move to next paragraph
                            # (Limitation: Only 1 formula per paragraph supported in MVP)
                            break
                        except Exception as e:
                            print(f"Error inserting OMML: {e}")
                            fixes.append(
                                {
                                    "id": f"err_latex_{para_idx}",
                                    "rule_id": "latex_to_omml",
                                    "description": f"Failed to insert formula: {e}",
                                    "status": "failed",
                                }
                            )

        return fixes


def convert_latex_in_document(doc: Document) -> list:
    """
    Convenience function to process LaTeX in a document.

    Args:
        doc: python-docx Document object

    Returns:
        List of fixes applied
    """
    converter = LaTeXConverter()
    return converter.process_document(doc)
